# -*- coding: utf-8 -*-
"""Genera la documentación Word del API Documento Soporte."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent


def setup_doc(title: str, subtitle: str = "Helpharma - API Documento Soporte") -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Versión del manual: 1.0  |  Proyecto: Ap-HelpharmaDocumentoSoporte")
    doc.add_paragraph("Confidencial - Uso interno Helpharma")
    doc.add_page_break()
    return doc


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def p(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph("")


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def save(doc: Document, filename: str):
    path = OUTPUT_DIR / filename
    doc.save(str(path))
    print(f"  Creado: {path.name}")


def doc_indice():
    doc = setup_doc("Índice General de Documentación", "Manual técnico y funcional")
    h1(doc, "Documentos incluidos en este manual")
    table(doc, ["#", "Archivo", "Contenido"], [
        ["01", "01_Funcionalidad_General.doc", "Descripción funcional y flujos principales"],
        ["02", "02_Proposito_y_Alcance.doc", "Para qué sirve, objetivos y alcance"],
        ["03", "03_Ficha_Tecnica.doc", "Stack tecnológico, dependencias y configuración"],
        ["04", "04_API_Endpoints.doc", "Endpoints, requests, responses y ejemplos de consumo"],
        ["05", "05_Arquitectura.doc", "Capas, patrones, DI y diagramas"],
        ["06", "06_Capa_Datos.doc", "Acceso a datos, SQL, bases de datos Ofima y DW"],
        ["07", "07_Seguridad.doc", "Autenticación API Key, CORS y buenas prácticas"],
        ["08", "08_Reglas_de_Negocio.doc", "Parseo de soportes, cuota moderadora, orígenes"],
        ["09", "09_Instalacion_Despliegue.doc", "Ejecución, debug, publicación y Swagger"],
        ["10", "10_Modelo_DTOs.doc", "Estructuras de datos de entrada y salida"],
    ])
    h1(doc, "Audiencia recomendada")
    bullets(doc, [
        "Desarrolladores .NET que mantengan o extiendan la API.",
        "Analistas funcionales que consuman los endpoints.",
        "Equipo de infraestructura / DevOps para despliegue.",
        "Integradores de sistemas externos (ERP, apps de entrega, etc.).",
    ])
    save(doc, "00_Indice_General.doc")


def doc_funcionalidad():
    doc = setup_doc("01 - Funcionalidad General")
    h1(doc, "1. Resumen")
    p(doc, "La API Documento Soporte es un servicio REST desarrollado en ASP.NET Core que "
          "consulta información de entregas farmacéuticas desde dos sistemas origen: "
          "Ofima (base HELPHARMA) y Data Warehouse / InfoPharma (base InfoPharma). "
          "Consolida datos de paciente, factura, medicamentos y valores de cuota moderadora "
          "para ser utilizados en la generación de documentos de soporte de entrega.")
    h1(doc, "2. Funcionalidades principales")
    numbered(doc, [
        "Consulta unificada de soporte por código compuesto (Ofima o DW).",
        "Consulta directa de soporte Ofima por tipo y número de documento (Trade).",
        "Consulta extendida de datos de soporte para entrega (DatosSoportes).",
        "Enrutamiento automático según formato del código de soporte.",
        "Exclusión de producto S3501 en consultas Ofima.",
        "Cálculo de cuota moderadora (producto S3500) con valor absoluto y acumulado.",
        "Protección de endpoints mediante API Key.",
        "Documentación interactiva con Swagger UI.",
    ])
    h1(doc, "3. Flujo funcional general")
    p(doc, "El consumidor envía un código de soporte o parámetros Trade. "
          "La capa de negocio determina el origen (Ofima o DW), ejecuta la consulta "
          "correspondiente en SQL Server y retorna un DTO estructurado en JSON.")
    h2(doc, "3.1 Flujo endpoint soportes/dctoprv")
    bullets(doc, [
        "Entrada: SoporteDto con campo Soporte.",
        "Business evalúa carácter en posición 2 del código.",
        "Ofima: Tipodcto (2 chars) + Nrodcto → PacienteDto.",
        "DW: Prefijo (3 chars) + NoEntrega (numérico) → PacienteDto.",
    ])
    h2(doc, "3.2 Flujo endpoint soportes/Trade")
    bullets(doc, [
        "Entrada: TradeDto (Tipodcto, Nrodcto).",
        "Consulta exclusivamente Ofima.",
        "Retorna PacienteDto con factura y líneas de medicamentos.",
    ])
    h2(doc, "3.3 Flujo endpoint soportes/DatosSoportes")
    bullets(doc, [
        "Entrada: SoporteDto; el controller elimina guiones del código.",
        "Business enruta a Ofima o DW igual que dctoprv.",
        "Retorna SoporteEntregaDto con datos ampliados (convenio, sede, plan, medicamentos).",
    ])
    h1(doc, "4. Sistemas origen")
    table(doc, ["Origen", "Base de datos", "Servidor (config)", "Tablas principales"], [
        ["Ofima", "HELPHARMA", "192.168.1.22", "trade, mvtrade, trademas, mtprocli, CANAL, TIPOVTA, MTBODEGA"],
        ["DW / InfoPharma", "InfoPharma", "10.0.0.5", "Entregas, MvEntregas, Pacientes, Convenios, Carteras, TiposEntrega"],
    ])
    save(doc, "01_Funcionalidad_General.doc")


def doc_proposito():
    doc = setup_doc("02 - Propósito y Alcance")
    h1(doc, "1. Propósito del sistema")
    p(doc, "Centralizar el acceso a la información de soportes de entrega de medicamentos "
          "dispersa en Ofima e InfoPharma, exponiéndola mediante una API HTTP segura y "
          "estandarizada para aplicaciones consumidoras que generan o validan documentos de soporte.")
    h1(doc, "2. Problema que resuelve")
    bullets(doc, [
        "Evita que cada aplicación consumidora se conecte directamente a múltiples bases SQL.",
        "Unifica reglas de negocio (cuota moderadora, filtro S3501, parseo de códigos).",
        "Ofrece contratos JSON consistentes (PacienteDto, SoporteEntregaDto).",
        "Permite integración controlada con autenticación por API Key.",
    ])
    h1(doc, "3. Usuarios / consumidores esperados")
    bullets(doc, [
        "Aplicaciones web o móviles de soporte de entrega.",
        "Servicios de generación de PDF / documentos.",
        "Integraciones internas Helpharma (ERP, logística, facturación).",
        "Equipos de soporte que consultan datos de paciente y entrega.",
    ])
    h1(doc, "4. Alcance incluido")
    bullets(doc, [
        "Consulta de datos de paciente, factura/entrega y medicamentos.",
        "Tres endpoints POST documentados en Swagger.",
        "Conexión de solo lectura a bases Ofima e InfoPharma.",
    ])
    h1(doc, "5. Alcance excluido")
    bullets(doc, [
        "No persiste ni modifica datos en las bases origen.",
        "No genera archivos PDF ni documentos físicos.",
        "No implementa autenticación de usuarios finales (solo API Key).",
        "No incluye caché ni colas de mensajería.",
    ])
    h1(doc, "6. Objetivos de negocio")
    numbered(doc, [
        "Disponibilizar información confiable de soportes en tiempo real.",
        "Reducir duplicidad de lógica SQL en aplicaciones cliente.",
        "Facilitar auditoría y trazabilidad mediante API centralizada.",
    ])
    save(doc, "02_Proposito_y_Alcance.doc")


def doc_ficha():
    doc = setup_doc("03 - Ficha Técnica")
    h1(doc, "1. Identificación del producto")
    table(doc, ["Campo", "Valor"], [
        ["Nombre", "API Documento Soporte (WebApp)"],
        ["Repositorio", "Ap-HelpharmaDocumentoSoporte"],
        ["Solución", "WebApp.slnx"],
        ["Framework", ".NET 10.0 (net10.0)"],
        ["Tipo", "ASP.NET Core Web API"],
        ["Puerto desarrollo HTTP", "5200"],
        ["Puerto desarrollo HTTPS", "7135"],
    ])
    h1(doc, "2. Estructura de proyectos")
    table(doc, ["Proyecto", "Tipo", "Responsabilidad"], [
        ["WebApp", "ASP.NET Core Web", "Controllers, Middleware, Program.cs, Swagger"],
        ["Business", "Class Library", "Lógica de negocio, enrutamiento Ofima/DW"],
        ["Data", "Class Library", "Acceso SQL, Connection Factories"],
        ["Models", "Class Library", "DTOs request/response"],
    ])
    h1(doc, "3. Dependencias NuGet principales")
    table(doc, ["Paquete", "Versión", "Proyecto", "Uso"], [
        ["Swashbuckle.AspNetCore", "6.5.0", "WebApp", "Swagger / OpenAPI"],
        ["Microsoft.EntityFrameworkCore", "10.0.3", "WebApp", "Referenciado (no usado activamente en acceso datos)"],
        ["Microsoft.Data.SqlClient", "6.1.4", "Data", "Conexión SQL Server"],
        ["AutoMapper", "16.0.0", "Data", "Referenciado (no usado en código actual)"],
    ])
    h1(doc, "4. Configuración (appsettings.json)")
    p(doc, "Secciones principales:")
    bullets(doc, [
        "ConnectionStrings:DefaultConnectionOfima → HELPHARMA",
        "ConnectionStrings:DefaultConnectionDW → InfoPharma",
        "ApiSecurity:ApiKey → clave para header X-API-KEY",
        "Logging, AllowedHosts",
    ])
    p(doc, "IMPORTANTE: Las credenciales deben externalizarse en producción "
          "(variables de entorno, Azure Key Vault, User Secrets). No commitear secretos.")
    h1(doc, "5. Requisitos de ejecución")
    bullets(doc, [
        "SDK .NET 10.0",
        "Acceso de red a servidores SQL 192.168.1.22 y 10.0.0.5",
        "Usuario SQL con permisos de lectura (Usuarioconsulta)",
        "Windows / IIS o Kestrel para hosting",
    ])
    save(doc, "03_Ficha_Tecnica.doc")


def doc_endpoints():
    doc = setup_doc("04 - API Endpoints: Consumo y Respuestas")
    h1(doc, "1. Información general")
    table(doc, ["Aspecto", "Detalle"], [
        ["Base URL (dev)", "http://localhost:5200"],
        ["Prefijo API", "/api/DocSoporte"],
        ["Autenticación", "Header X-API-KEY"],
        ["Formato", "application/json"],
        ["Swagger", "http://localhost:5200/swagger/index.html"],
    ])
    h1(doc, "2. Endpoint: POST /api/DocSoporte/soportes/dctoprv")
    h2(doc, "Descripción")
    p(doc, "Obtiene información completa del paciente y factura a partir de un código de soporte unificado.")
    h2(doc, "Request body")
    code(doc, '{\n  "Soporte": "FE123456"\n}')
    h2(doc, "Response 200 - PacienteDto")
    code(doc, '{\n  "nombrePaciente": "...",\n  "tipoId": "CC",\n  "paciente": "1234567890",\n  "direccionPaciente": "...",\n  "telefonoPaciente": "...",\n  "celularPaciente": "...",\n  "complemento": "...",\n  "valorCMTotal": 15000.00,\n  "factura": {\n    "convenio": "...",\n    "fecha": "2026-03-15T00:00:00",\n    "bodega": "...",\n    "tipoEntrega": "...",\n    "cartera": "...",\n    "observacion": "...",\n    "usuario": "...",\n    "ordenes": [\n      {\n        "ordenes": "...",\n        "producto": "...",\n        "nombre": "...",\n        "cantidad": 1,\n        "lote": "...",\n        "valorMx": 0\n      }\n    ]\n  }\n}')
    h2(doc, "Errores")
    bullets(doc, ["400: datos obligatorios inválidos", "401: API Key faltante o inválida", "500: excepción (incluye stack trace en texto)"])
    h1(doc, "3. Endpoint: POST /api/DocSoporte/soportes/Trade")
    h2(doc, "Request body")
    code(doc, '{\n  "Tipodcto": "FE",\n  "Nrodcto": "123456"\n}')
    p(doc, "Consulta directamente Ofima. Response: mismo PacienteDto que dctoprv.")
    h1(doc, "4. Endpoint: POST /api/DocSoporte/soportes/DatosSoportes")
    h2(doc, "Descripción")
    p(doc, "Retorna datos extendidos para documento de soporte de entrega. Elimina guiones del código antes de procesar.")
    h2(doc, "Request body")
    code(doc, '{\n  "Soporte": "IFEMI-403592"\n}')
    h2(doc, "Response 200 - SoporteEntregaDto")
    code(doc, '{\n  "idConvenio": "...",\n  "nombreConvenio": "...",\n  "fecha": "2026-03-15T00:00:00",\n  "idBodega": "...",\n  "nombreSede": "...",\n  "nombreActividad": "",\n  "tipoEntrega": "...",\n  "tipoPlan": "...",\n  "idCartera": "...",\n  "nombrePaciente": "...",\n  "idTipoId": "CC",\n  "idPaciente": "1234567890",\n  "celular": "...",\n  "telefono": "...",\n  "direccion": "...",\n  "complemento": "...",\n  "observacion": "...",\n  "valorCM": "15000",\n  "medicamentos": [ ... ]\n}')
    h2(doc, "Errores")
    bullets(doc, ['400: { "mensaje": "..." } — incluye errores SQL de conversión'])
    h1(doc, "5. Ejemplos de consumo")
    h2(doc, "cURL")
    code(doc, 'curl -X POST "http://localhost:5200/api/DocSoporte/soportes/DatosSoportes" ^\n  -H "Content-Type: application/json" ^\n  -H "X-API-KEY: SU_API_KEY" ^\n  -d "{\\"Soporte\\": \\"IFEMI-403592\\"}"')
    h2(doc, "PowerShell")
    code(doc, '$headers = @{ "X-API-KEY" = "SU_API_KEY"; "Content-Type" = "application/json" }\n$body = \'{ "Soporte": "IFEMI-403592" }\'\nInvoke-RestMethod -Uri "http://localhost:5200/api/DocSoporte/soportes/DatosSoportes" `\n  -Method POST -Headers $headers -Body $body')
    h2(doc, "C# HttpClient")
    code(doc, 'var client = new HttpClient();\nclient.DefaultRequestHeaders.Add("X-API-KEY", "SU_API_KEY");\nvar json = JsonSerializer.Serialize(new { Soporte = "IFEMI-403592" });\nvar content = new StringContent(json, Encoding.UTF8, "application/json");\nvar response = await client.PostAsync(\n  "http://localhost:5200/api/DocSoporte/soportes/DatosSoportes", content);')
    h1(doc, "6. Códigos HTTP")
    table(doc, ["Código", "Significado"], [
        ["200", "Consulta exitosa"],
        ["400", "Validación o error de negocio/SQL (DatosSoportes)"],
        ["401", "API Key requerida o inválida"],
        ["500", "Error interno (dctoprv, Trade) o ApiKey no configurada"],
    ])
    save(doc, "04_API_Endpoints.doc")


def doc_arquitectura():
    doc = setup_doc("05 - Arquitectura del Sistema")
    h1(doc, "1. Patrón arquitectónico")
    p(doc, "Arquitectura en capas (N-Tier) con separación de responsabilidades:")
    bullets(doc, [
        "Presentación: WebApp (Controllers, Middleware, Pipeline HTTP)",
        "Negocio: Business (DocSuportBusiness, reglas de enrutamiento)",
        "Acceso a datos: Data (SQL directo, factories de conexión)",
        "Modelos: Models (DTOs desacoplados)",
    ])
    h1(doc, "2. Diagrama de capas")
    code(doc, """
┌─────────────────────────────────────────────────────────┐
│                    CLIENTE HTTP                          │
│              (Swagger, App, Postman, cURL)                 │
└─────────────────────────┬───────────────────────────────┘
                          │ X-API-KEY
                          ▼
┌─────────────────────────────────────────────────────────┐
│  WebApp - Pipeline                                       │
│  Swagger → ApiKeyMiddleware → Controllers                │
└─────────────────────────┬───────────────────────────────┘
                          ▼
┌─────────────────────────────────────────────────────────┐
│  Business - DocSuportBusiness                            │
│  Enrutamiento Ofima / DW según formato soporte           │
└───────────────┬─────────────────────┬───────────────────┘
                ▼                     ▼
┌───────────────────────┐  ┌──────────────────────────────┐
│ DocumentoSoporte      │  │ DocumentoSoporteDWData       │
│ OfimaData             │  │                              │
└───────────┬───────────┘  └──────────────┬───────────────┘
            ▼                             ▼
┌───────────────────────┐  ┌──────────────────────────────┐
│ OfimaConnectionFactory│  │ DwConnectionFactory          │
│ HELPHARMA             │  │ InfoPharma                   │
└───────────────────────┘  └──────────────────────────────┘
""")
    h1(doc, "3. Inyección de dependencias")
    p(doc, "Registro en DependencyContainer.cs y DataAccessDependency.cs:")
    table(doc, ["Interfaz", "Implementación", "Lifetime"], [
        ["IDocSoportBusiness", "DocSuportBusiness", "Scoped"],
        ["IDocumentoSoporteOfimaData", "DocumentoSoporteOfimaData", "Scoped"],
        ["IDocumentoSoporteDWData", "DocumentoSoporteDWData", "Scoped"],
        ["IOfimaConnectionFactory", "OfimaConnectionFactory", "Scoped"],
        ["IDwConnectionFactory", "DwConnectionFactory", "Scoped"],
    ])
    h1(doc, "4. Pipeline HTTP (Program.cs)")
    numbered(doc, [
        "UseSwagger / UseSwaggerUI",
        "UseWhen → ApiKeyMiddleware (excepto /swagger)",
        "UseHttpsRedirection",
        "UseAuthorization",
        "MapControllers",
    ])
    h1(doc, "5. Interfaces clave")
    bullets(doc, [
        "IDocSoportBusiness: contrato de negocio",
        "IDocumentoSoporteOfimaData / IDocumentoSoporteDWData: contratos de datos",
        "IOfimaConnectionFactory / IDwConnectionFactory: creación de SqlConnection",
    ])
    save(doc, "05_Arquitectura.doc")


def doc_datos():
    doc = setup_doc("06 - Capa de Datos y Bases de Datos")
    h1(doc, "1. Estrategia de acceso")
    p(doc, "Acceso mediante ADO.NET (Microsoft.Data.SqlClient) con consultas SQL parametrizadas. "
          "No se utiliza Entity Framework para las operaciones actuales a pesar de estar referenciado.")
    h1(doc, "2. Connection Factories")
    h2(doc, "OfimaConnectionFactory")
    p(doc, "Lee ConnectionStrings:DefaultConnectionOfima y crea SqlConnection a HELPHARMA.")
    h2(doc, "DwConnectionFactory")
    p(doc, "Lee ConnectionStrings:DefaultConnectionDW y crea SqlConnection a InfoPharma.")
    h1(doc, "3. DocumentoSoporteOfimaData - Métodos")
    table(doc, ["Método", "Parámetros", "Retorno", "Descripción"], [
        ["GetSoporteOfima", "DCTOPRV", "PacienteDto", "Por código DCTOPRV (legacy)"],
        ["GetSoporteTradeOfima", "TradeDto", "PacienteDto", "Por TIPODCTO + NRODCTO"],
        ["GetDatosSoportes", "TradeDto", "SoporteEntregaDto", "Datos extendidos entrega"],
    ])
    h2(doc, "Tablas Ofima utilizadas")
    p(doc, "trade, mvtrade, trademas, CANAL, mtprocli, tipocar, TIPOVTA, MTBODEGA")
    h2(doc, "Filtro común Ofima")
    p(doc, "m.PRODUCTO <> 'S3501' — excluye línea de producto S3501.")
    h1(doc, "4. DocumentoSoporteDWData - Métodos")
    table(doc, ["Método", "Parámetros", "Retorno", "Descripción"], [
        ["GetSoporteDW", "prefijo, noEntrega (int)", "PacienteDto", "Entrega por prefijo y número"],
        ["GetDatosSoportes", "TradeDto", "SoporteEntregaDto", "Datos extendidos DW"],
    ])
    h2(doc, "Tablas DW utilizadas")
    p(doc, "MvEntregas, Entregas, Convenios, Pacientes, Carteras, TiposEntrega")
    h2(doc, "Hint NOLOCK")
    p(doc, "Consultas DW usan WITH (NOLOCK) para lectura sin bloqueo.")
    h1(doc, "5. Mapeo reader → DTO")
    p(doc, "Cada fila del SqlDataReader se mapea manualmente a OrdenDto. "
          "La primera fila inicializa datos de cabecera (paciente/entrega); "
          "filas subsiguientes agregan medicamentos a la lista.")
    save(doc, "06_Capa_Datos.doc")


def doc_seguridad():
    doc = setup_doc("07 - Seguridad y Configuración")
    h1(doc, "1. Autenticación API Key")
    p(doc, "Middleware ApiKeyMiddleware valida header X-API-KEY en todas las rutas excepto /swagger.")
    table(doc, ["Escenario", "HTTP", "Mensaje"], [
        ["Sin header", "401", "Api Key requerida."],
        ["Key incorrecta", "401", "Api Key inválida."],
        ["Key no configurada en servidor", "500", "ApiKey no configurada en el servidor."],
        ["Swagger UI", "Sin validación", "Acceso libre a documentación"],
    ])
    h1(doc, "2. Configuración de la clave")
    p(doc, "appsettings.json → ApiSecurity:ApiKey. En producción usar variables de entorno:")
    code(doc, "ApiSecurity__ApiKey=clave_segura_produccion")
    h1(doc, "3. CORS")
    p(doc, "Política AllowLocalhost registrada en Program.cs con AllowAnyOrigin. "
          "Nota: verificar que UseCors esté aplicado en pipeline si se requiere consumo cross-origin desde browser.")
    h1(doc, "4. HTTPS")
    p(doc, "UseHttpsRedirection activo. En desarrollo disponible http://localhost:5200 y https://localhost:7135.")
    h1(doc, "5. Recomendaciones de seguridad")
    numbered(doc, [
        "No almacenar credenciales SQL ni API Key en el repositorio Git.",
        "Rotar API Key periódicamente.",
        "Restringir Swagger en producción o proteger con autenticación adicional.",
        "No retornar ex.ToString() al cliente en producción (filtra stack traces).",
        "Usar cuenta SQL de solo lectura con permisos mínimos.",
        "Considerar rate limiting y logging de auditoría.",
    ])
    save(doc, "07_Seguridad.doc")


def doc_reglas():
    doc = setup_doc("08 - Reglas de Negocio")
    h1(doc, "1. Determinación de origen (Ofima vs DW)")
    p(doc, "Regla en DocSuportBusiness: se evalúa el carácter en índice 2 (tercer carácter) del código de soporte.")
    table(doc, ["Condición", "Origen", "Parseo Tipodcto/Nrodcto (DatosSoportes)"], [
        ["char.IsDigit(soporte[2]) == true", "Ofima", "Tipodcto = 2 primeros chars; Nrodcto = resto"],
        ["char.IsDigit(soporte[2]) == false", "DW", "Tipodcto = 3 primeros chars; Nrodcto = resto (string en SQL)"],
    ])
    h2(doc, "Ejemplos")
    table(doc, ["Soporte", "Origen", "Tipodcto", "Nrodcto"], [
        ["FE123456", "Ofima", "FE", "123456"],
        ["ABC1234567", "DW", "ABC", "1234567"],
        ["IFEMI403592 (sin guion)", "DW", "IFE", "MI403592 — ERROR SQL si NoEntrega es int"],
    ])
    h1(doc, "2. Cuota moderadora (producto S3500)")
    bullets(doc, [
        "Si producto == S3500, ValorMx / ValorCM se convierte a valor absoluto (Math.Abs).",
        "ValorCMTotal (PacienteDto) o ValorCM (SoporteEntregaDto) acumula solo líneas S3500.",
        "Producto S3501 se excluye de consultas Ofima.",
    ])
    h1(doc, "3. Normalización en DatosSoportes")
    p(doc, "El controller elimina guiones: Soporte.Replace(\"-\", \"\"). "
          "Ejemplo: IFEMI-403592 → IFEMI403592.")
    h1(doc, "4. GetSoporte (dctoprv) - diferencia con DatosSoportes")
    p(doc, "En ruta DW de GetSoporte se usa int.Parse(soporte.Substring(3)) para noEntrega. "
          "En DatosSoportes el Nrodcto se pasa como string al SQL.")
    h1(doc, "5. Casos límite conocidos")
    bullets(doc, [
        "Códigos con prefijo alfanumérico largo (ej. IFEMI) pueden parsearse incorrectamente con regla de 3 caracteres.",
        "NIT de Ofima mapeado a IdPaciente en SoporteEntregaDto es string (documento identidad).",
        "Soporte vacío o muy corto puede causar IndexOutOfRangeException.",
    ])
    save(doc, "08_Reglas_de_Negocio.doc")


def doc_instalacion():
    doc = setup_doc("09 - Instalación, Despliegue y Depuración")
    h1(doc, "1. Requisitos previos")
    bullets(doc, [
        ".NET SDK 10.0",
        "Visual Studio 2022 / VS Code / Cursor con extensión C#",
        "Acceso VPN/red a servidores SQL",
    ])
    h1(doc, "2. Compilación")
    code(doc, 'dotnet build "api\\WebApp\\WebApp\\WebApp.csproj"')
    h1(doc, "3. Ejecución en desarrollo")
    code(doc, 'dotnet run --project "api\\WebApp\\WebApp\\WebApp.csproj" --launch-profile http')
    p(doc, "URL: http://localhost:5200/swagger/index.html")
    h1(doc, "4. Depuración (F5)")
    p(doc, "Configuración en .vscode/launch.json → perfil WebApp Debug. "
          "Breakpoints recomendados: DocSoporteController líneas 74-75, "
          "DocSuportBusiness líneas 57 y 72-77, DocumentoSoporteDWData línea 185.")
    h1(doc, "5. Publicación")
    p(doc, "Perfil FolderProfile disponible en Properties/PublishProfiles. "
          "Publicar a carpeta o IIS según infraestructura Helpharma.")
    h1(doc, "6. Verificación post-despliegue")
    numbered(doc, [
        "Acceder a /swagger (si está habilitado).",
        "Probar endpoint con X-API-KEY válida.",
        "Verificar conectividad a ambas bases SQL.",
        "Revisar logs de aplicación.",
    ])
    h1(doc, "7. Detener la aplicación")
    code(doc, "netstat -ano | findstr :5200\nStop-Process -Id <PID> -Force")
    save(doc, "09_Instalacion_Despliegue.doc")


def doc_dtos():
    doc = setup_doc("10 - Modelo de Datos (DTOs)")
    h1(doc, "1. DTOs de entrada (Request)")
    h2(doc, "SoporteDto")
    table(doc, ["Campo", "Tipo", "Validación", "Descripción"], [
        ["Soporte", "string", "Required", "Código compuesto de soporte"],
    ])
    h2(doc, "TradeDto")
    table(doc, ["Campo", "Tipo", "Validación", "Descripción"], [
        ["Tipodcto", "string", "Required", "Tipo documento Ofima (2 chars) o prefijo DW (3 chars)"],
        ["Nrodcto", "string", "Required", "Número documento o entrega"],
    ])
    h1(doc, "2. PacienteDto (respuesta dctoprv / Trade)")
    table(doc, ["Campo", "Tipo", "Descripción"], [
        ["NombrePaciente", "string", "Nombre completo"],
        ["TipoId", "string", "Tipo identificación"],
        ["Paciente", "string", "Número identificación / NIT"],
        ["DireccionPaciente", "string", "Dirección"],
        ["TelefonoPaciente", "string", "Teléfono fijo"],
        ["CelularPaciente", "string", "Celular"],
        ["Complemento", "string", "Complemento dirección"],
        ["ValorCMTotal", "decimal", "Total cuota moderadora acumulada"],
        ["Factura", "FacturaDto", "Datos de factura y líneas"],
    ])
    h1(doc, "3. FacturaDto")
    table(doc, ["Campo", "Tipo"], [
        ["Convenio", "string"], ["Fecha", "DateTime"], ["Bodega", "string"],
        ["TipoEntrega", "string"], ["Cartera", "string"], ["Observacion", "string"],
        ["Usuario", "string"], ["Ordenes", "List<OrdenDto>"],
    ])
    h1(doc, "4. OrdenDto (línea medicamento)")
    table(doc, ["Campo", "Tipo"], [
        ["Ordenes", "string"], ["Producto", "string"], ["Nombre", "string"],
        ["Cantidad", "int"], ["Lote", "string"], ["ValorMx", "decimal"],
    ])
    h1(doc, "5. SoporteEntregaDto (respuesta DatosSoportes)")
    table(doc, ["Campo", "Tipo", "Descripción"], [
        ["IdConvenio", "string", "Código convenio"],
        ["NombreConvenio", "string", "Nombre convenio"],
        ["Fecha", "DateTime", "Fecha entrega"],
        ["IdBodega", "string", "Código bodega"],
        ["NombreSede", "string", "Nombre sede / farmacia"],
        ["NombreActividad", "string", "Actividad (vacío en consulta actual)"],
        ["TipoEntrega", "string", "Tipo de entrega"],
        ["TipoPlan", "string", "Tipo plan"],
        ["IdCartera", "string", "Código cartera"],
        ["NombrePaciente", "string", "Nombre paciente"],
        ["idTipoId", "string", "Tipo identificación"],
        ["IdPaciente", "string", "Identificación paciente"],
        ["Celular", "string", "Celular"],
        ["Telefono", "string", "Teléfono"],
        ["Direccion", "string", "Dirección"],
        ["Complemento", "string", "Complemento"],
        ["Observacion", "string", "Observaciones"],
        ["ValorCM", "string", "Valor cuota moderadora formateado"],
        ["medicamentos", "List<OrdenDto>", "Lista de medicamentos entregados"],
    ])
    save(doc, "10_Modelo_DTOs.doc")


def main():
    print("Generando documentación en:", OUTPUT_DIR)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc_indice()
    doc_funcionalidad()
    doc_proposito()
    doc_ficha()
    doc_endpoints()
    doc_arquitectura()
    doc_datos()
    doc_seguridad()
    doc_reglas()
    doc_instalacion()
    doc_dtos()
    print("\nDocumentación generada exitosamente (11 archivos .doc).")


if __name__ == "__main__":
    main()
